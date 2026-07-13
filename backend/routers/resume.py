from fastapi import APIRouter, UploadFile, File, HTTPException, Depends
from sqlalchemy.orm import Session
from typing import List
from database import get_db
from models import ResumeAnalysis, AiActivity, AiActivityResponse
from services.claude_service import (
    analyze_resume,
    generate_resume_bullets_generic,
    create_interview_questions_generic,
)
from services.ai_errors import raise_clean_ai_error
from auth import Owner, get_owner, owned, log_activity
from pydantic import BaseModel
import pypdf
import io
import json
import logging

logger = logging.getLogger(__name__)

router = APIRouter()


class ResumeTextRequest(BaseModel):
    resume_text: str


def extract_pdf_text(data: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(data))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages).strip()


def extract_docx_text(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs).strip()


@router.post("/analyze")
async def analyze_resume_file(
    file: UploadFile = File(...),
    owner: Owner = Depends(get_owner),
    db: Session = Depends(get_db),
):
    content = await file.read()
    filename = file.filename or ""
    lower = filename.lower()

    if lower.endswith(".doc") and not lower.endswith(".docx"):
        raise HTTPException(
            status_code=400,
            detail="The legacy .doc format isn't supported. Please save your resume as .docx or PDF and try again.",
        )

    try:
        if lower.endswith(".pdf"):
            text = extract_pdf_text(content)
        elif lower.endswith(".docx"):
            text = extract_docx_text(content)
        elif lower.endswith(".txt"):
            text = content.decode("utf-8", errors="ignore")
        else:
            raise HTTPException(
                status_code=400,
                detail="Unsupported format. Upload a PDF, DOCX, or TXT file.",
            )
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(
            status_code=400,
            detail="Could not read this file. Make sure it's a valid PDF, DOCX, or TXT resume.",
        )

    if len(text.strip()) < 50:
        raise HTTPException(
            status_code=422,
            detail="Could not extract readable text. Ensure the file is not image-only.",
        )

    try:
        analysis = await analyze_resume(text)
    except Exception as e:
        raise_clean_ai_error(logger, "Resume analysis", e)

    row = ResumeAnalysis(
        filename=filename,
        resume_text=text,
        ats_score=analysis.get("ats_score", 0),
        analysis_json=json.dumps(analysis),
        guest_id=owner.guest_id,
        user_id=owner.user_id,
    )
    db.add(row)
    db.commit()
    db.refresh(row)

    seeker_document_id = None
    try:
        from services.seeker_document_storage import save_seeker_bytes, extension_allowed
        from models import SeekerDocument
        if extension_allowed(filename):
            owner_key = str(owner.user_id or owner.guest_id or "anon")
            path, mime, size, digest = save_seeker_bytes(owner_key, filename, content)
            prev = (
                owned(db.query(SeekerDocument), SeekerDocument, owner)
                .filter(SeekerDocument.document_type == "resume", SeekerDocument.deleted_at.is_(None))
                .count()
            )
            doc = SeekerDocument(
                user_id=owner.user_id,
                guest_id=owner.guest_id,
                document_type="resume",
                file_name=filename,
                mime_type=mime,
                file_size=size,
                version_number=prev + 1,
                storage_path=path,
                source_resume_analysis_id=row.id,
                content_sha256=digest,
            )
            db.add(doc)
            db.commit()
            db.refresh(doc)
            seeker_document_id = doc.id
    except Exception:
        logger.exception("Failed to retain resume bytes for upload assist (non-fatal)")

    log_activity(
        db, owner, "resume_analyzed",
        f"Analyzed resume — ATS Score: {analysis.get('ats_score', 0)}%",
        filename,
    )

    return {
        "filename": filename,
        "resume_text": text,
        "analysis": analysis,
        "resume_id": row.id,
        "seeker_document_id": seeker_document_id,
    }


@router.get("/history", response_model=List[dict])
async def resume_history(
    owner: Owner = Depends(get_owner),
    db: Session = Depends(get_db),
):
    rows = (
        owned(db.query(ResumeAnalysis), ResumeAnalysis, owner)
        .order_by(ResumeAnalysis.created_at.desc())
        .limit(20)
        .all()
    )
    return [
        {
            "id": r.id,
            "filename": r.filename,
            "resume_text": r.resume_text,
            "ats_score": r.ats_score,
            "analysis": json.loads(r.analysis_json) if r.analysis_json else None,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in rows
    ]


@router.post("/analyze-text")
async def analyze_resume_text(body: dict):
    text = body.get("text", "")
    if len(text.strip()) < 50:
        raise HTTPException(status_code=422, detail="Resume text is too short.")
    try:
        analysis = await analyze_resume(text)
    except Exception as e:
        raise_clean_ai_error(logger, "Resume analysis", e)
    return {"analysis": analysis}


@router.post("/bullets")
async def get_resume_bullets(
    request: ResumeTextRequest,
    owner: Owner = Depends(get_owner),
    db: Session = Depends(get_db),
):
    if len(request.resume_text.strip()) < 50:
        raise HTTPException(status_code=422, detail="Resume text is too short.")
    try:
        bullets = await generate_resume_bullets_generic(request.resume_text)
    except Exception as e:
        raise_clean_ai_error(logger, "Bullet generation", e)
    log_activity(db, owner, "bullets_generated", f"Generated {len(bullets)} improved resume bullets")
    return {"bullets": bullets}


@router.post("/interview-questions")
async def get_interview_questions(
    request: ResumeTextRequest,
    owner: Owner = Depends(get_owner),
    db: Session = Depends(get_db),
):
    if len(request.resume_text.strip()) < 50:
        raise HTTPException(status_code=422, detail="Resume text is too short.")
    try:
        questions = await create_interview_questions_generic(request.resume_text)
    except Exception as e:
        raise_clean_ai_error(logger, "Interview question generation", e)
    log_activity(db, owner, "questions_generated", f"Generated {len(questions)} interview questions")
    return {"questions": questions}
